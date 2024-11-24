from PyPDF2 import PdfReader
import os
from crewai import Agent, Task, Crew
from openai import OpenAI
from typing import Optional
from dotenv import load_dotenv
import datetime
import time
import docx


class AARProcessor:
    def __init__(self, openai_api_key: Optional[str] = None):
        load_dotenv()
        
        # Initialize OpenAI client with Omnistack configuration
        self.client = OpenAI(
            base_url="https://api.omnistack.sh/openai/v1",
            api_key="osk_f6ffd2937fdacb98de5a3c37e128e89d"
        )

        # Initialize agents with Omnistack configuration
        self.llm_config = {
            "client": self.client,
            "model": "jonathan_simone_carmen"
        }

        # Initialize agents
        self.document_processor = self._create_document_processor()
        self.translator = self._create_translator()
        self.summarizer = self._create_summarizer()
        self.quality_checker = self._create_quality_checker()

    def _create_document_processor(self):
        return Agent(
            role="Document Processor",
            goal="Extract and structure content from Ukrainian AAR documents",
            backstory="You are an expert in processing military documents, "
            "particularly After-Action Reports. You understand the "
            "structure and importance of military documentation and "
            "can effectively extract all relevant information from various document formats.",
            allow_delegation=False,
            verbose=True,
            llm_config=self.llm_config
        )

    def _create_translator(self):
        return Agent(
            role="Military Translator",
            goal="Accurately translate Ukrainian military content to English",
            backstory="You are a specialized military translator with deep "
            "knowledge of both Ukrainian and English military terminology. "
            "You understand military slang, idioms, and cultural nuances "
            "specific to Ukrainian armed forces. Your translations maintain "
            "the precise meaning and context of the original text.",
            allow_delegation=False,
            verbose=True,
            llm_config=self.llm_config
        )

    def _create_summarizer(self):
        return Agent(
            role="Military Intelligence Summarizer",
            goal="Create concise, actionable summaries of translated AARs",
            backstory="You are an intelligence analyst specialized in creating "
            "clear, actionable summaries from military reports. You "
            "understand what information is most critical for frontline "
            "soldiers and can present it in a clear, structured format "
            "that enables quick decision-making.",
            allow_delegation=False,
            verbose=True,
            llm_config=self.llm_config
        )

    def _create_quality_checker(self):
        return Agent(
            role="Quality Assurance Specialist",
            goal="Ensure accuracy and completeness of translations and summaries",
            backstory="You are a quality control expert with experience in "
            "military documentation. You verify that translations "
            "maintain accuracy, context, and military terminology "
            "while ensuring summaries contain all critical information "
            "in an accessible format.",
            allow_delegation=False,
            verbose=True,
            llm_config=self.llm_config
        )

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text content from a document file (PDF, DOC, or DOCX)."""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()

            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")

        except Exception as e:
            raise Exception(f"Error extracting text from file: {str(e)}")

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOC/DOCX file."""
        doc = docx.Document(file_path)  
        paragraphs = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n".join(paragraphs)

    def process_aar(self, file_path: str):
        timings = []
        start_time = time.time()

        # Extract text from document
        try:
            extraction_start = time.time()
            text_content = self.extract_text_from_file(file_path)
            extraction_time = time.time() - extraction_start
            timings.append(
                (f"Document Extraction ({os.path.splitext(file_path)[1]})", extraction_time))
            print(f"\nExtracted text (first 200 chars):\n{text_content[:200]}\n")
            print(f"Total length of extracted text: {len(text_content)} characters\n")
        except Exception as e:
            print(f"Error reading document: {str(e)}")
            return None, None

        # Create tasks for the workflow
        tasks = []

        # Task 1: Document Processing
        task_start = time.time()
        task1 = Task(
            description="Process and structure the extracted document content",
            agent=self.document_processor,
            context=[{
                "description": "Process the following text extracted from the AAR document",
                "input": f"Process the following text extracted from the AAR document "
                f"and structure it appropriately:\n\n{text_content}",
                "expected_output": "A structured representation of the AAR content"
            }],
            expected_output="A structured representation of the AAR content with clear sections and formatting"
        )
        tasks.append(task1)

        # Task 2: Translation
        task2 = Task(
            description="Translate the structured content from Ukrainian to English",
            agent=self.translator,
            context=[{
                "description": "Translate the processed content",
                # This will be replaced with actual output
                "input": "{{previous_task.output}}",
                "expected_output": "An accurate English translation"
            }],
            expected_output="An accurate English translation of the AAR that maintains military terminology and context"
        )
        tasks.append(task2)

        # Task 3: Summarization
        task3 = Task(
            description="Create a concise, actionable summary",
            agent=self.summarizer,
            context=[{
                "description": "Generate a clear summary",
                # This will be replaced with actual output
                "input": "{{previous_task.output}}",
                "expected_output": "A concise, actionable summary"
            }],
            expected_output="A concise, actionable summary of the AAR highlighting key insights and recommendations"
        )
        tasks.append(task3)

        # Task 4: Quality Check
        task4 = Task(
            description="Verify translation and summary quality",
            agent=self.quality_checker,
            context=[{
                "description": "Review the translation and summary",
                # This will be replaced with actual output
                "input": "{{previous_task.output}}",
                "expected_output": "A quality assessment report"
            }],
            expected_output="A confirmation that the translation and summary meet quality standards"
        )
        tasks.append(task4)

        # Create and run the crew
        crew_start_time = time.time()
        crew = Crew(
            agents=[self.document_processor, self.translator,
                    self.summarizer],
            tasks=tasks,
            verbose=True
        )

        # Track individual task times
        task_times = {}

        def task_callback(task, output):
            task_times[task.description] = time.time()
            return output

        # Initialize start time for first task
        task_times["start"] = time.time()

        # Run the crew with task tracking
        result = crew.kickoff()

        # Calculate individual task durations
        task_descriptions = [t.description for t in tasks]
        for i in range(len(task_descriptions)):
            current_task = task_descriptions[i]
            next_time = task_times.get(
                task_descriptions[i + 1], time.time()) if i < len(task_descriptions) - 1 else time.time()
            duration = next_time - \
                (task_times.get(
                    task_descriptions[i-1], task_times["start"]) if i > 0 else task_times["start"])
            timings.append((f"Task {i+1}: {current_task}", duration))

        crew_time = time.time() - crew_start_time
        timings.append(("Total Crew Time", crew_time))

        total_time = time.time() - start_time
        timings.append(("Total Processing", total_time))

        return result, timings


def main():
    # Example usage
    processor = AARProcessor()

    # Process all files in docs directory
    docs_dir = "docs"
    supported_extensions = ['.pdf', '.doc', '.docx']

    if not os.path.exists(docs_dir):
        print(f"Error: Directory '{docs_dir}' does not exist!")
        return

    # Get all supported files in the docs directory
    files_to_process = []
    for file in os.listdir(docs_dir):
        if any(file.lower().endswith(ext) for ext in supported_extensions):
            files_to_process.append(os.path.join(docs_dir, file))

    if not files_to_process:
        print(f"No supported files found in '{docs_dir}' directory!")
        print(f"Supported file types: {', '.join(supported_extensions)}")
        return

    print(f"\nFound {len(files_to_process)} files to process:")
    for file in files_to_process:
        print(f"- {os.path.basename(file)}")
    print()

    # Process each file
    for file_path in files_to_process:
        print(f"\nProcessing: {os.path.basename(file_path)}")
        print("-" * 50)

        result, timings = processor.process_aar(file_path)

        if result is None:
            print(f"Failed to process {os.path.basename(file_path)}")
            continue

        # Save result to markdown file
        output_dir = "output"
        os.makedirs(output_dir, exist_ok=True)

        # Generate timestamp for unique filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = os.path.splitext(os.path.basename(file_path))[0]
        output_file = os.path.join(output_dir, f"aar_analysis_{base_filename}_{timestamp}.md")
        timing_file = os.path.join(output_dir, "processing_times.md")

        # Save analysis results
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("# After-Action Report Analysis\n\n")
            f.write(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Source file: {os.path.basename(file_path)}\n\n")
            f.write("## Analysis Results\n\n")
            f.write(str(result))  # Convert CrewOutput to string

        # Save or append timing results
        if os.path.exists(timing_file):
            with open(timing_file, "r", encoding="utf-8") as f:
                existing_content = f.read()
        else:
            existing_content = "# Processing Times Log\n\n"

        with open(timing_file, "w", encoding="utf-8") as f:
            f.write(existing_content)
            f.write(f"\n## Run at {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {os.path.basename(file_path)}\n\n")
            f.write("| Step | Time (seconds) |\n")
            f.write("|------|----------------|\n")
            for step, duration in timings:
                f.write(f"| {step} | {duration:.2f} |\n")
            f.write("\n")

        print(f"\nAnalysis saved to: {output_file}")
        print(f"Timing information appended to: {timing_file}")
        print("-" * 50)


if __name__ == "__main__":
    main()
